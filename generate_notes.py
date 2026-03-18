import docx

doc = docx.Document()
doc.add_heading('Skill Gap Analyzer - Project Notes', 0)

doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph("Skill Gap Analyzer is a full-stack web application designed to help university students identify the gap between their current programming/technical skill set and the actual skills required for their target job roles in the industry. By analyzing a student's selected branch, current skills, and desired job role, the system calculates a match percentage and highlights the missing skills they need to focus on.")

doc.add_heading('2. Technology Stack', level=1)
doc.add_paragraph("• Backend: Python, Django Framework\n• Frontend: HTML5, CSS3, Bootstrap 5, JavaScript\n• Database: SQLite\n• Authentication: Django built-in authentication system")

doc.add_heading('3. Key Features', level=1)
doc.add_paragraph("1. User Authentication: Secure login, registration, and logout functionality.\n2. Interactive Analysis Tool: A multi-step form allowing users to seamlessly select their academic Branch, followed by their existing Skills, and finally their target Job Role.\n3. Gap Analysis Engine: Compares the user's selected skills with the industry-standard skills required for the target role to generate a precise Match Percentage and a list of Missing Skills.\n4. Dynamic Data Loading: Smooth UI that filters available skills and roles based on the selected academic branch.\n5. Custom Admin Panel: Built-in Django administrative dashboard to manage the lists of branches, skills, and roles.")

doc.add_heading('4. Database Architecture (Models)', level=1)
doc.add_paragraph("• Branch: Represents academic departments (e.g., Computer Science, Mechanical Engineering).\n• Skill: Represents standalone technical skills (e.g., Python, AutoCAD, Data Structures).\n• JobRole: Relevant job titles linked to specific Branches (e.g., Software Engineer).\n• RequiredSkill: A mapping table connecting Job Roles with their necessary Skills.\n• UserProfile: Extends the default User model to store additional information like college name and branch.\n• StudentSkillSelection: Stores the results of an analysis session, keeping a history of match percentages.")

doc.add_heading('5. Local Setup & Execution', level=1)
doc.add_paragraph("1. Open your terminal in the project folder.\n2. Activate the virtual environment: venv\\Scripts\\activate\n3. Install project dependencies: pip install -r requirements.txt\n4. Apply database migrations: python manage.py migrate\n5. Start the development server: python manage.py runserver\n6. Access the website in a browser at http://127.0.0.1:8000/")

doc.save('Skill_Gap_Analyzer_Notes.docx')
print("Notes document generated as Skill_Gap_Analyzer_Notes.docx")
